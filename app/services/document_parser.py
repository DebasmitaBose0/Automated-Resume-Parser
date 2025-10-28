import pdfplumber
from docx import Document
import os
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """Service for extracting text from PDF and Word documents."""
    
    @staticmethod
    def extract_text_from_pdf(file_path):
        """
        Extract text from PDF file using pdfplumber.
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            str: Extracted text content
        """
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path):
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_doc(file_path):
        """
        Extract text from DOC file. Note: This is a simplified implementation.
        For production use, consider using python-docx2txt or similar libraries.
        
        Args:
            file_path (str): Path to the DOC file
            
        Returns:
            str: Extracted text content
        """
        try:
            # For .doc files, we'll try to use python-docx anyway
            # In production, you might want to use a more robust solution
            return DocumentParser.extract_text_from_docx(file_path)
        except Exception as e:
            logger.error(f"Error extracting text from DOC {file_path}: {str(e)}")
            # Fallback: try to read as plain text (not recommended for production)
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    return file.read()
            except Exception as fallback_error:
                logger.error(f"Fallback extraction failed: {str(fallback_error)}")
                raise Exception(f"Failed to extract text from DOC: {str(e)}")
    
    @classmethod
    def extract_text(cls, file_path):
        """
        Extract text from a document based on its file extension.
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            str: Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return cls.extract_text_from_docx(file_path)
        elif file_extension == '.doc':
            return cls.extract_text_from_doc(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    @staticmethod
    def clean_text(text):
        """
        Clean and normalize extracted text.
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace and normalize line breaks
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        cleaned_text = '\n'.join(lines)
        
        # Remove multiple consecutive spaces
        import re
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        
        return cleaned_text.strip()